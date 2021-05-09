import openpyxl
import pprint
import docx,os
from docx.shared import Pt


doc=docx.Document()
# doc.add_heading('The North Eastern institute')
dictData={}
wb=openpyxl.load_workbook('ReadingExcelAndUpdatingWord.xlsx')
sheet=wb.get_sheet_by_name('Names')
for i in range(sheet.max_row):
	if i==0:
		continue
	dictData[sheet.cell(row=i+1,column=1).value]=sheet.cell(row=i+1,column=2).value


for key,val in dictData.items():
	prepara=doc.add_paragraph('The North Eastern institute')
	prepara.style="Heading 1"
	prepara.alignment=1
	prepara.runs[0].add_break()
	firstpara=doc.add_paragraph('')
	first_run=firstpara.add_run('Received an amount of Rs.'+ str(val) +' from ')
	second_run=firstpara.add_run(str(key))
	second_run.bold=True
	first_run.font.size=Pt(14)
	second_run.font.size=Pt(14)
	second_run.font.name='Times New Roman'
	first_run.font.name='Times New Roman'

	second_run.add_break()
	secondpara=doc.add_paragraph(key)
	secondpara.alignment=2
	secondpara.runs[0].font.size=Pt(24)
	secondpara.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

	

doc.save("sample3.docx")
os.system('start sample3.docx')





# wb.save('produceSales.xlsx')