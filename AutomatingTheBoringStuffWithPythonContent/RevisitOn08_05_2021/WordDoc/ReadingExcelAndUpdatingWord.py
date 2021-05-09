import openpyxl
import pprint
import docx,os


doc=docx.Document()
doc.add_heading('The REAL meaning of the universe')
dictData={}
wb=openpyxl.load_workbook('ReadingExcelAndUpdatingWord.xlsx')
sheet=wb.get_sheet_by_name('Names')
for i in range(sheet.max_row):
	if i==0:
		continue
	dictData[sheet.cell(row=i+1,column=1).value]=sheet.cell(row=i+1,column=2).value

for key,val in dictData.items():
	print(key+":"+str(val))
	firstpara=doc.add_paragraph('')
	first_run=firstpara.add_run(key+" :")
	second_run=firstpara.add_run(str(val))
	second_run.underline=True
	first_run.bold=True
	second_run.add_break(docx.enum.text.WD_BREAK.PAGE)

	firstpara.style="Intense Quote"

doc.save("sample.docx")
os.system('start sample.docx')





# wb.save('produceSales.xlsx')