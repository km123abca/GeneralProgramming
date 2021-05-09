import openpyxl
import pprint
import docx,os


doc=docx.Document()
doc.add_heading('The REAL meaning of the universe')
dictData={}
wb=openpyxl.load_workbook('ReadingExcelAndUpdatingWord.xlsx')
sheet=wb.get_sheet_by_name('Names')
for i in range(sheet.max_row):
	if i==0:
		continue
	dictData[sheet.cell(row=i+1,column=1).value]=sheet.cell(row=i+1,column=2).value


for key,val in dictData.items():
	table = doc.add_table(rows=1, cols=2)
	table.cell(0, 0).text=key
	table.cell(0, 1).text=str(val)
	table.style = 'LightShading-Accent1'
	firstpara=doc.add_paragraph('')
	first_run=firstpara.add_run('')
	first_run.add_break(docx.enum.text.WD_BREAK.PAGE)
	# first_run=firstpara.add_run(key+" :")
	# second_run=firstpara.add_run(str(val))
	# second_run.underline=True
	# first_run.bold=True
	# second_run.add_break(docx.enum.text.WD_BREAK.PAGE)

	# firstpara.style="Intense Quote"

doc.save("sample2.docx")
os.system('start sample2.docx')





# wb.save('produceSales.xlsx')