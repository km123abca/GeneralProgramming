import docx,os
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


def aligntext(para,num):
	for p in para:
		p.alignment=num

ffile=open('Files\\SomeNames.txt')


doc=docx.Document()

####### THIS IS WHERE WE ADD OUR CUSTOM STYLES
obj_styles = doc.styles
obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(20)
obj_font.name = 'Brush Script MT'

obj_charstyle2 = obj_styles.add_style('CommentsStyle2', WD_STYLE_TYPE.CHARACTER)
obj_font2 = obj_charstyle2.font
obj_font2.size = Pt(20)
obj_font2.name = 'Times New Roman'


linelist=['It would be a pleasure to have the company of','Robocop','at 11010 Memory Lane on the Evening of','April 1st','at 7 o\' clock']
stylelist=['CommentsStyle','CommentsStyle2','CommentsStyle','CommentsStyle2','CommentsStyle']
paragraph1=None
for name in ffile.readlines():
	count=0
	for line in linelist:
		paragraph1=doc.add_paragraph('')
		if count==1:
			paragraph1.add_run(name[:-1],style=stylelist[count])
		else:
			paragraph1.add_run(line,style=stylelist[count])
		if count in [1,3]:
			paragraph1.runs[0].bold=True
		else:
			paragraph1.runs[0].italic=True
		paragraph1.runs[0].add_break()
		count+=1
	paragraph1.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)


aligntext(doc.paragraphs,1)

ffile.close()
doc.save('Files\\invitation.docx')
os.system('start Files\\invitation.docx')