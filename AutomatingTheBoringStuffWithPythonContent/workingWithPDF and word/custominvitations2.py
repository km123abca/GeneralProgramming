import docx,os
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


def aligntext(para,num):
	for p in para:
		p.alignment=num

doc=docx.Document()

####### THIS IS WHERE WE ADD OUR CUSTOM STYLES
obj_styles = doc.styles
obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(20)
obj_font.name = 'Brush Script MT'

obj_charstyle2 = obj_styles.add_style('CommentsStyle2', WD_STYLE_TYPE.CHARACTER)
obj_font2 = obj_charstyle2.font
obj_font2.size = Pt(20)
obj_font2.name = 'Times New Roman'

# OUR PARAGRAPHS ARE WRITTEN INDIVIDUALLY HERE
paragraph1=doc.add_paragraph('')
paragraph1.add_run('It would be a pleasure to have the company of ', style = 'CommentsStyle')
doc.paragraphs[0].runs[0].italic=True
doc.paragraphs[0].runs[0].add_break()


paragraph1=doc.add_paragraph('')
paragraph1.add_run('Robocop ', style = 'CommentsStyle2')
doc.paragraphs[1].runs[0].bold=True
doc.paragraphs[1].runs[0].italic=False
doc.paragraphs[1].runs[0].add_break()


paragraph1=doc.add_paragraph('')
paragraph1.add_run('at the event ', style = 'CommentsStyle')
doc.paragraphs[2].runs[0].italic=True
aligntext(doc.paragraphs,1)


doc.save('Files\\invitation.docx')
os.system('start Files\\invitation.docx')