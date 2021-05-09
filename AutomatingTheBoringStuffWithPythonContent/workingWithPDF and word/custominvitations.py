import docx,os
from docx.shared import Pt

#print(doc.paragraphs[0].runs[0].font.size)



def fontsize(fname,fsize,docname):
	stylel=docname.styles['Normal']
	fontl=stylel.font
	fontl.name=fname
	fontl.size=Pt(fsize)

def aligntext(para,num):
	for p in para:
		p.alignment=num

doc=docx.Document()
'''
obj_styles = document.styles
obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(10)
obj_font.name = 'Times New Roman'
'''
#fontsize('Brush Script MT',20,doc)
doc.add_paragraph('It would be a pleasure to have the company of ')
#doc.paragraphs[0].style=doc.styles['Normal']
doc.paragraphs[0].runs[0].italic=True
doc.paragraphs[0].runs[0].add_break()

#doc.paragraphs[0].alignment = 3 # for left, 1 for right, 2 center, 3 justify ....

#fontsize('Times New Roman',20,doc)
doc.add_paragraph('Robocop ')
doc.paragraphs[1].runs[0].bold=True
doc.paragraphs[1].runs[0].italic=False
doc.paragraphs[1].runs[0].add_break()


fontsize('Brush Script MT',20,doc)
doc.add_paragraph('at the event')
doc.paragraphs[2].runs[0].italic=True

aligntext(doc.paragraphs,1)


#doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
#doc.add_paragraph('I hope we meet again')

doc.save('Files\\invitation.docx')
os.system('start Files\\invitation.docx')